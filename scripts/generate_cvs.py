from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("output/cv")
ACCENT = RGBColor(31, 77, 120)
TEXT = RGBColor(28, 28, 28)
MUTED = RGBColor(80, 80, 80)


CONTACT = {
    "name": "Saheed (Prosper) Tobi Ogundipe",
    "mobile_title": "Mobile Engineer | React Native | Frontend Lead",
    "web_title": "Frontend Lead | Full-Stack Engineer | React, Next.js, Node.js",
    "email": "prosper8890@gmail.com",
    "phone": "+234 903 402 3685 | +234 701 771 8795",
    "location": "Lagos, Nigeria - Available for relocation",
    "linkedin": "linkedin.com/in/oluwatobi3685",
    "github": "github.com/prosper-85",
    "portfolio": "prosper-01.netlify.app",
}


EXPERIENCE = [
    {
        "role": "Senior Frontend Engineer / Frontend Lead",
        "company": "iRecharge Tech Innovations LTD",
        "dates": "Jan 2024 - Present",
        "mobile": [
            "Lead frontend delivery for fintech products, including merchant and consumer payment experiences across web and mobile-adjacent flows.",
            "Contributed to Blink Merchant, a business payment platform supporting tap-to-pay, geo-payment, and scan-to-pay flows.",
            "Contributed to Blink Pay, a virtual card application focused on secure personal transactions.",
            "Improved platform performance, UI consistency, and communication between product, backend, and frontend teams.",
        ],
        "web": [
            "Lead frontend engineering for fintech platforms including iRecharge and Accelerate, with ownership across architecture, UI consistency, and delivery quality.",
            "Built scalable React and Next.js UI systems that improved development speed and reduced repeated implementation work.",
            "Collaborated with backend and product teams to shape API contracts, improve data handling, and deliver production-ready workflows.",
            "Improved performance and user experience across payment and business-facing platforms.",
        ],
    },
    {
        "role": "Frontend / Full-Stack Engineer",
        "company": "Orion Industrial Tech",
        "dates": "Nov 2022 - Dec 2023",
        "mobile": [
            "Developed and shipped cross-platform mobile features with React Native for iOS and Android.",
            "Contributed to Dryve User App, a ride-hailing platform with real-time location and user interaction flows.",
            "Reduced unnecessary re-renders and improved mobile load behavior through component and state optimization.",
            "Integrated backend APIs and collaborated on mobile-friendly contracts for reliable data handling.",
        ],
        "web": [
            "Built web and mobile application features with React, React Native, and API-backed data flows.",
            "Optimized rendering behavior by reducing avoidable re-renders and improving load-time performance.",
            "Integrated backend APIs and strengthened frontend data handling across product screens.",
            "Partnered with backend engineers to define practical API contracts for production interfaces.",
        ],
    },
    {
        "role": "Frontend / Full-Stack Engineer",
        "company": "Spacepen Technology",
        "dates": "Apr 2022",
        "mobile": [
            "Built responsive product interfaces and backend integrations for payroll automation workflows.",
            "Implemented dynamic UI flows for employee management and transaction tracking.",
            "Collaborated on frontend-backend integration to support stable internal business tools.",
        ],
        "web": [
            "Developed React and Next.js interfaces for payroll automation systems.",
            "Implemented dynamic workflows for employee management and transaction tracking.",
            "Delivered responsive, performance-optimized interfaces for internal business users.",
        ],
    },
    {
        "role": "Frontend Developer",
        "company": "FICLC",
        "dates": "Feb 2020",
        "mobile": [
            "Developed React and Next.js interfaces for payroll automation and internal operations.",
            "Implemented reusable UI flows and integrated backend services for business users.",
        ],
        "web": [
            "Developed user interfaces with React and Next.js for payroll automation systems.",
            "Implemented dynamic UI workflows for employee management and transaction tracking.",
            "Contributed to backend API integration for seamless frontend-backend interaction.",
        ],
    },
]


PROJECTS = {
    "mobile": [
        {
            "name": "Blink Merchant",
            "meta": "Digital payment platform for businesses - Play Store: play.google.com/store/apps/details?id=com.blinkbusiness",
            "points": [
                "Business payment product supporting tap-to-pay, geo-payment, and scan-to-pay transaction flows.",
                "Relevant areas: secure payment UX, merchant workflows, mobile fintech interaction patterns.",
            ],
        },
        {
            "name": "Blink Pay",
            "meta": "Virtual card application - Play Store: play.google.com/store/apps/details?id=com.blinkpay",
            "points": [
                "Consumer payment app for secure personal transactions and virtual card usage.",
                "Relevant areas: transaction confidence, account flows, mobile financial UI states.",
            ],
        },
        {
            "name": "Dryve User App",
            "meta": "Ride-hailing platform - Play Store: play.google.com/store/apps/details?id=com.dryveuser",
            "points": [
                "Ride-hailing user app with real-time location, trip flow, and user interaction features.",
                "Relevant areas: location-driven UX, state synchronization, mobile interaction reliability.",
            ],
        },
        {
            "name": "Sanitrack",
            "meta": "React, React Native, TypeScript, Express.js, MongoDB, Docker",
            "points": [
                "Designed and implemented web and mobile features with consistent cross-platform interfaces.",
                "Built backend services and structured API response handling to prevent frontend data inconsistencies.",
            ],
        },
    ],
    "web": [
        {
            "name": "Accelerate",
            "meta": "Lead Frontend contribution - accelerate.ng",
            "points": [
                "Contributed frontend leadership to business-facing fintech workflows at iRecharge.",
                "Improved UI consistency, delivery quality, and collaboration with backend/product stakeholders.",
            ],
        },
        {
            "name": "iRecharge",
            "meta": "Lead Frontend contribution - irecharge.ng",
            "points": [
                "Worked across payment platform interfaces, performance, and production user experience.",
                "Supported reusable frontend patterns for payment and recharge workflows.",
            ],
        },
        {
            "name": "Sanitrack",
            "meta": "React, React Native, TypeScript, Express.js, MongoDB, Docker",
            "points": [
                "Designed full-stack web and mobile features with secure API communication and structured validation flows.",
                "Applied error handling and response modeling to reduce frontend data inconsistencies.",
            ],
        },
        {
            "name": "Payroll Automation Interfaces",
            "meta": "React, Next.js, REST APIs",
            "points": [
                "Built dynamic workflows for employee management and transaction tracking.",
                "Delivered responsive and performance-optimized interfaces for internal business users.",
            ],
        },
    ],
}


SKILLS = {
    "mobile": [
        "Mobile: React Native, Expo, Firebase Auth, Firestore, navigation, cross-platform UI patterns, app performance",
        "Frontend: React, Next.js, TypeScript, reusable components, design systems, responsive UI",
        "State and Forms: Redux Toolkit, Context API, React Hook Form, Zod, state synchronization",
        "Backend Collaboration: Node.js, Express, NestJS, REST APIs, JWT auth, MongoDB, PostgreSQL",
        "Delivery: Docker, GitHub, Sentry, ESLint, production debugging, performance optimization",
    ],
    "web": [
        "Frontend Architecture: React, Next.js, TypeScript strict mode, component-driven architecture, reusable UI libraries",
        "UI Engineering: design systems, responsive interfaces, accessibility awareness, Tailwind CSS, performance optimization",
        "State and Forms: Redux Toolkit, Context API, React Hook Form, Zod, API integration, data modeling",
        "Backend and APIs: Node.js, Express, NestJS, REST APIs, JWT auth, MongoDB, PostgreSQL",
        "Leadership and Delivery: frontend leadership, cross-team communication, code quality, Sentry, ESLint, GitHub",
    ],
}


SUMMARY = {
    "mobile": (
        "Mobile-focused Frontend Lead and Full-Stack Engineer with 5+ years of experience building fintech, mobility, "
        "and enterprise products. Strong in React Native, Expo, React, Next.js, TypeScript, Firebase, and API-backed "
        "mobile workflows. Experienced with payment products, ride-hailing flows, real-time location experiences, "
        "performance optimization, and cross-team delivery. Available for relocation for local or international roles."
    ),
    "web": (
        "Frontend Lead and Full-Stack Engineer with 5+ years of experience building scalable fintech, mobility, and "
        "enterprise applications. Specialized in React, Next.js, TypeScript, reusable UI systems, API integration, and "
        "backend collaboration with Node.js, Express, and NestJS. Known for communication, leadership, and production-ready "
        "delivery. Available for relocation for local or international roles."
    ),
}


def set_run_font(run, size=9, bold=None, color=TEXT, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=3, line_spacing=1.05):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_hyperlink(paragraph, text, url, size=8.4):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4D78")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run_element.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.06

    for style_name, size in [("Heading 1", 10.8), ("Heading 2", 10.2)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.keep_with_next = True


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_header(doc, variant):
    title = CONTACT["mobile_title" if variant == "mobile" else "web_title"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    run = p.add_run(CONTACT["name"])
    set_run_font(run, size=17, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    contact_text = f"{CONTACT['email']} | {CONTACT['phone']} | {CONTACT['location']}"
    run = p.add_run(contact_text)
    set_run_font(run, size=9.1, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=5)
    add_hyperlink(p, CONTACT["linkedin"], f"https://{CONTACT['linkedin']}", size=9.1)
    run = p.add_run(" | ")
    set_run_font(run, size=9.1, color=MUTED)
    add_hyperlink(p, CONTACT["github"], f"https://{CONTACT['github']}", size=9.1)
    run = p.add_run(" | ")
    set_run_font(run, size=9.1, color=MUTED)
    add_hyperlink(p, CONTACT["portfolio"], f"https://{CONTACT['portfolio']}", size=9.1)
    add_rule(p)


def add_section(doc, title):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(title.upper())
    return paragraph


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4, line_spacing=1.1)
    run = p.add_run(text)
    set_run_font(run, size=9.6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_spacing(p, after=1.6, line_spacing=1.06)
    run = p.add_run(text)
    set_run_font(run, size=9)
    return p


def add_role(doc, role, company, dates, bullets):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0.5)
    run = p.add_run(role)
    set_run_font(run, size=9.7, bold=True, color=RGBColor(0, 0, 0))
    run = p.add_run(f" | {company} | {dates}")
    set_run_font(run, size=9.1, color=MUTED)
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_project(doc, project):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0.5)
    run = p.add_run(project["name"])
    set_run_font(run, size=9.7, bold=True, color=RGBColor(0, 0, 0))
    run = p.add_run(f" | {project['meta']}")
    set_run_font(run, size=8.9, color=MUTED)
    for point in project["points"]:
        add_bullet(doc, point)


def add_footer(section, variant):
    footer = section.footer.paragraphs[0]
    footer.text = ""


def build_cv(variant):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.36)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.18)
    configure_styles(doc)
    add_footer(section, variant)

    add_header(doc, variant)

    add_section(doc, "Professional Summary")
    add_body_paragraph(doc, SUMMARY[variant])

    add_section(doc, "Core Skills")
    for skill in SKILLS[variant]:
        add_bullet(doc, skill)

    add_section(doc, "Professional Experience")
    for job in EXPERIENCE:
        add_role(doc, job["role"], job["company"], job["dates"], job[variant])

    add_section(doc, "Selected Projects")
    for project in PROJECTS[variant]:
        add_project(doc, project)

    add_section(doc, "Education, Certifications and Languages")
    add_bullet(doc, "Bachelor of Statistics - University of Ilorin, 2021")
    add_bullet(doc, "React & TypeScript - Udemy, 2022")
    add_bullet(doc, "React Native - Udemy, 2022")
    add_bullet(doc, "English; Yoruba")

    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    files = {
        "mobile": OUTPUT_DIR / "Saheed_Prosper_Tobi_Mobile_Engineer_CV.docx",
        "web": OUTPUT_DIR / "Saheed_Prosper_Tobi_Web_Frontend_Engineer_CV.docx",
    }
    for variant, path in files.items():
        doc = build_cv(variant)
        doc.core_properties.author = CONTACT["name"]
        doc.core_properties.title = (
            "Saheed Prosper Tobi - Mobile Engineer CV"
            if variant == "mobile"
            else "Saheed Prosper Tobi - Web Frontend Engineer CV"
        )
        doc.save(path)
        print(path)


if __name__ == "__main__":
    main()
